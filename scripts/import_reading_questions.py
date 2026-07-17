from __future__ import annotations

import argparse
import html
import json
import re
import sqlite3
import string
from collections import defaultdict
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


QUESTION_START = re.compile(r"^\s*(\d{1,2})\s*[.．]\s*(.*)$", re.S)
ANALYSIS_START = re.compile(r"^\s*(\d{1,2})\s*[.．]\s*([（(].*)$", re.S)
OPTION_LABEL = re.compile(r"(?m)(?:^|\n)\s*([A-D])(?:\s*[.．、]|\s+)|(?<![A-Za-z])([A-D])\s*[.．、]")
ANSWER_LABEL = "【参考答案】"
PRACTICAL_LABEL = "【实战解析】"
NOTE_LABEL = "花生批注"
TYPE_LABEL = re.compile(r"【题型[^】]*】\s*([^\n]+)")

KNOWN_OPTION_REPAIRS: dict[tuple[int, int], dict[str, str]] = {
    (4, 19): {"A": "⑤③①②④⑥", "B": "⑤①③⑥④②", "C": "①③④⑥②⑤", "D": "①③②④⑥⑤"},
    (9, 20): {"A": "⑥④①③⑤②", "B": "⑤④①③⑥②", "C": "⑤⑥②①④③", "D": "⑥②⑤④①③"},
    (14, 18): {"A": "⑥⑦①⑤②③④", "B": "④⑦①②⑤③⑥", "C": "⑥⑦③①④⑤②", "D": "④⑦③①⑤②⑥"},
}

CIRCLED = (
    [chr(value) for value in range(0x2460, 0x2474)]
    + [chr(value) for value in range(0x3251, 0x3260)]
    + [chr(value) for value in range(0x32B1, 0x32C0)]
)
CIRCLED_VALUE = {char: index + 1 for index, char in enumerate(CIRCLED)}
MARKER_CHARS = set(CIRCLED + list(string.digits))

NOISE_PATTERNS = (
    re.compile(r"^\s*$"),
    re.compile(r"^练习题\s*\d+\s*套\s*$"),
    re.compile(r"^练习题\s*\d+\s*$"),
    re.compile(r"^片段阅读\s*6?0?0?\s*[（(]?[上下]册[）)]?\s*$"),
    re.compile(r"^四海公考\s*$"),
    re.compile(r"^SIHAI\s*GONG\s*KAO\s*$", re.I),
    re.compile(r"^\d{1,3}\s*$"),
)


@dataclass
class StyledChar:
    char: str
    bold: bool = False
    underline: bool = False
    italic: bool = False


def normalize_space(text: str) -> str:
    text = text.replace("\u3000", " ").replace("\xa0", " ").replace("\t", " ")
    text = re.sub(r"[ \r\f\v]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    return text.strip()


def compact_text(text: str) -> str:
    text = normalize_space(text)
    text = re.sub(r"(?<=[\u3400-\u9fff，。；：！？、（）《》“”‘’]) +(?=[\u3400-\u9fff，。；：！？、（）《》“”‘’])", "", text)
    text = re.sub(r"(?<=[\u3400-\u9fff]) +(?=[A-D](?:选项|项))", "", text)
    text = re.sub(r"(?<=[A-D]) +(?=(?:选项|项))", "", text)
    return text.strip()


def canonical_labels(text: str) -> str:
    text = re.sub(r"【\s*参考答案\s*】", ANSWER_LABEL, text)
    text = re.sub(r"【\s*实战解析\s*】", PRACTICAL_LABEL, text)
    text = re.sub(r"【\s*(题型[^】]*?)\s*】", lambda match: f"【{compact_text(match.group(1))}】", text)
    return text


def match_text(text: str) -> str:
    return re.sub(r"[^0-9A-Za-z\u3400-\u9fff]", "", compact_text(text)).lower()


def is_noise(text: str) -> bool:
    text = normalize_space(text)
    return any(pattern.fullmatch(text) for pattern in NOISE_PATTERNS)


def leaf_body_paragraphs(document: DocumentObject) -> list[Paragraph]:
    result: list[Paragraph] = []
    for element in document.element.body.iter():
        if element.tag != qn("w:p"):
            continue
        if any(child.tag == qn("w:p") for child in element.iterdescendants()):
            continue
        result.append(Paragraph(element, document))
    return result


def all_body_paragraphs(document: DocumentObject) -> list[Paragraph]:
    return [Paragraph(element, document) for element in document.element.body.iter() if element.tag == qn("w:p")]


def textbox_streams(document: DocumentObject) -> list[list[Paragraph]]:
    streams: list[list[Paragraph]] = []
    for box in document.element.body.iter():
        if box.tag != qn("w:txbxContent"):
            continue
        paragraphs = [Paragraph(element, document) for element in box.iter() if element.tag == qn("w:p")]
        if paragraphs:
            streams.append(paragraphs)
    return streams


def paragraph_chars(paragraph: Paragraph) -> list[StyledChar]:
    chars: list[StyledChar] = []
    if paragraph.runs:
        for run in paragraph.runs:
            underline = bool(run.underline)
            for char in run.text:
                chars.append(StyledChar(char, bool(run.bold), underline, bool(run.italic)))
    else:
        chars.extend(StyledChar(char) for char in paragraph.text)
    return chars


def stream_chars(paragraphs: Iterable[Paragraph]) -> list[StyledChar]:
    result: list[StyledChar] = []
    kept = 0
    for paragraph in paragraphs:
        if is_noise(paragraph.text):
            continue
        paragraph_values = paragraph_chars(paragraph)
        paragraph_text = "".join(value.char for value in paragraph_values)
        ad_positions = [position for marker in ("公考最新资料", "更新进度微信") if (position := paragraph_text.find(marker)) >= 0]
        if ad_positions:
            paragraph_values = paragraph_values[: min(ad_positions)]
        if not paragraph_values:
            continue
        if kept:
            result.append(StyledChar("\n"))
        result.extend(paragraph_values)
        kept += 1
    return result


def normalize_styled_chars(chars: list[StyledChar]) -> list[StyledChar]:
    normalized: list[StyledChar] = []
    for item in chars:
        char = item.char.replace("\u3000", " ").replace("\t", " ").replace("\r", "\n")
        if char.isspace() and char != "\n":
            char = "\xa0" if item.underline else " "
        if char == " " and normalized and normalized[-1].char in {" ", "\n"}:
            continue
        if char == "\n" and normalized and normalized[-1].char == "\n":
            continue
        normalized.append(StyledChar(char, item.bold, item.underline, item.italic))

    def cjk_or_punctuation(value: str) -> bool:
        return bool(re.fullmatch(r"[\u3400-\u9fff，。；：！？、（）《》“”‘’]", value))

    cleaned: list[StyledChar] = []
    for index, item in enumerate(normalized):
        if item.char == " " and not item.underline:
            previous = normalized[index - 1].char if index else ""
            following = normalized[index + 1].char if index + 1 < len(normalized) else ""
            if cjk_or_punctuation(previous) and cjk_or_punctuation(following):
                continue
        cleaned.append(item)
    while cleaned and cleaned[0].char in {" ", "\n"}:
        cleaned.pop(0)
    while cleaned and cleaned[-1].char in {" ", "\n"}:
        cleaned.pop()
    return cleaned


def compress_runs(chars: list[StyledChar]) -> list[dict[str, object]]:
    chars = normalize_styled_chars(chars)
    runs: list[dict[str, object]] = []
    for item in chars:
        style = (item.bold, item.underline, item.italic)
        if runs and (bool(runs[-1].get("bold")), bool(runs[-1].get("underline")), bool(runs[-1].get("italic"))) == style:
            runs[-1]["text"] = str(runs[-1]["text"]) + item.char
            continue
        run: dict[str, object] = {"text": item.char}
        if item.bold:
            run["bold"] = True
        if item.underline:
            run["underline"] = True
        if item.italic:
            run["italic"] = True
        runs.append(run)
    return runs


def runs_text(runs: list[dict[str, object]]) -> str:
    return compact_text("".join(str(run["text"]).replace("\xa0", " ") for run in runs))


def slice_chars(chars: list[StyledChar], start: int, end: int | None = None) -> list[StyledChar]:
    return chars[start:end]


def find_option_spans(text: str, require_four: bool = True) -> list[tuple[str, int, int]]:
    matches = list(OPTION_LABEL.finditer(text))
    label = lambda item: item.group(1) or item.group(2)
    for start_index, match in enumerate(matches):
        if label(match) != "A":
            continue
        selected = [match]
        expected = "B"
        for candidate in matches[start_index + 1 :]:
            if label(candidate) == expected:
                selected.append(candidate)
                if expected == "D":
                    return [(label(item), item.start(), item.end()) for item in selected]
                expected = chr(ord(expected) + 1)
    if not require_four:
        unique: list[re.Match[str]] = []
        seen: set[str] = set()
        for match in matches:
            if label(match) not in seen:
                unique.append(match)
                seen.add(label(match))
        if unique:
            return [(label(item), item.start(), item.end()) for item in unique]
    return []


def parse_question_content(paragraphs: list[Paragraph], analysis_header: bool = False, allow_partial_options: bool = False) -> tuple[list[dict[str, object]], str, list[dict[str, object]], dict[str, list[dict[str, object]]]]:
    chars = stream_chars(paragraphs)
    text = "".join(item.char for item in chars)
    header_pattern = ANALYSIS_START if analysis_header else QUESTION_START
    header = header_pattern.match(text)
    if not header:
        raise ValueError(f"question chunk does not start with a numbered question: {text[:160]!r}")
    content_start = header.end(1)
    while content_start < len(text) and text[content_start] in " .．\n":
        content_start += 1
    if analysis_header:
        first_newline = text.find("\n", content_start)
        content_start = first_newline + 1 if first_newline >= 0 else content_start

    option_spans = find_option_spans(text[content_start:], require_four=not allow_partial_options)
    if len(option_spans) < (1 if allow_partial_options else 4):
        raise ValueError(f"expected four options, found {len(option_spans)}: {text[:500]!r}")
    option_spans = [(label, start + content_start, end + content_start) for label, start, end in option_spans]
    stem_runs = compress_runs(slice_chars(chars, content_start, option_spans[0][1]))
    option_runs: dict[str, list[dict[str, object]]] = {}
    options: list[dict[str, object]] = []
    for index, (label, _start, content) in enumerate(option_spans):
        end = option_spans[index + 1][1] if index + 1 < len(option_spans) else len(chars)
        runs = compress_runs(slice_chars(chars, content, end))
        plain = runs_text(runs)
        plain = re.split(r"【参考答案】|【题型|【实战解析】", plain, maxsplit=1)[0].strip()
        if index == 3 and plain != runs_text(runs):
            keep = len(plain)
            rebuilt: list[dict[str, object]] = []
            consumed = 0
            for run in runs:
                value = str(run["text"])
                if consumed >= keep:
                    break
                part = value[: keep - consumed]
                if part:
                    new_run = dict(run)
                    new_run["text"] = part
                    rebuilt.append(new_run)
                consumed += len(part)
            runs = rebuilt
        options.append({"label": label, "content": plain})
        option_runs[label] = runs
    return stem_runs, runs_text(stem_runs), options, option_runs


def split_numbered_chunks(paragraphs: list[Paragraph], analysis: bool) -> list[list[Paragraph]]:
    pattern = ANALYSIS_START if analysis else QUESTION_START
    starts: list[int] = []
    for index, paragraph in enumerate(paragraphs):
        text = normalize_space(paragraph.text)
        match = pattern.match(text)
        if match and 1 <= int(match.group(1)) <= 20:
            starts.append(index)
    return [paragraphs[start:end] for start, end in zip(starts, starts[1:] + [len(paragraphs)])]


def parse_book(path: Path, first_set: int) -> list[dict[str, object]]:
    document = Document(path)
    # Some option rows are stored in floating text boxes, so include descendant paragraphs.
    chunks = split_numbered_chunks(all_body_paragraphs(document), analysis=False)
    if len(chunks) != 300:
        raise ValueError(f"{path.name}: expected 300 book questions, found {len(chunks)}")
    questions: list[dict[str, object]] = []
    for index, chunk in enumerate(chunks):
        expected_local = index % 20 + 1
        actual = int(QUESTION_START.match(normalize_space(chunk[0].text)).group(1))
        if actual != expected_local:
            raise ValueError(f"{path.name}: question sequence mismatch at {index + 1}: {actual} != {expected_local}")
        stem_rich, stem, options, option_rich = parse_question_content(chunk, allow_partial_options=True)
        questions.append(
            {
                "setNumber": first_set + index // 20,
                "localNumber": expected_local,
                "stem": stem,
                "stemRich": stem_rich,
                "options": options,
                "optionRich": option_rich,
            }
        )
    return questions


def paragraph_plain_stream(paragraphs: list[Paragraph]) -> str:
    return "\n".join(canonical_labels(compact_text(paragraph.text)) for paragraph in paragraphs if paragraph.text and not is_noise(paragraph.text))


def extract_source(header: str) -> tuple[str, int | None, str]:
    header = normalize_space(header)
    match = ANALYSIS_START.match(header)
    value = match.group(2).strip("（）() ") if match else header
    accuracy_match = re.search(r"(\d{1,3})\s*%", value)
    accuracy = int(accuracy_match.group(1)) if accuracy_match else None
    source = re.sub(r"\s*\d{1,3}\s*%\s*$", "", value).strip()
    return source or "片段阅读600题", accuracy, value


def parse_analysis_candidate(chunk: list[Paragraph]) -> dict[str, object] | None:
    if not chunk:
        return None
    header = normalize_space(chunk[0].text)
    start = ANALYSIS_START.match(header)
    if not start:
        return None
    local_number = int(start.group(1))
    full_text = paragraph_plain_stream(chunk)
    answer_match = re.search(r"【参考答案\s*】\s*([A-Da-d])", full_text)
    type_match = TYPE_LABEL.search(full_text)

    practical_parts: list[str] = []
    practical_started = False
    for paragraph in chunk:
        text = canonical_labels(compact_text(paragraph.text))
        if not text or is_noise(text):
            continue
        if PRACTICAL_LABEL in text:
            practical_started = True
            text = text.split(PRACTICAL_LABEL, 1)[1].strip()
        if not practical_started:
            continue
        if NOTE_LABEL in text:
            before = text.split(NOTE_LABEL, 1)[0]
            before = re.sub(r"[①-⑳㉑-㊿0-9]+$", "", before).strip("：: ")
            if before:
                practical_parts.append(before)
            break
        if ANALYSIS_START.match(text) or ANSWER_LABEL in text or TYPE_LABEL.search(text):
            continue
        practical_parts.append(text)

    try:
        annotated_stem_rich, annotated_stem, annotated_options, annotated_option_rich = parse_question_content(chunk, analysis_header=True, allow_partial_options=True)
    except Exception:
        return None
    source, accuracy, source_label = extract_source(header)
    return {
        "localNumber": local_number,
        "source": source,
        "sourceLabel": source_label,
        "accuracy": accuracy,
        "answer": answer_match.group(1).upper() if answer_match else None,
        "typeAndPassage": compact_text(type_match.group(1)) if type_match else None,
        "practicalAnalysis": "\n".join(practical_parts).strip(),
        "annotatedStem": annotated_stem,
        "annotatedStemRich": annotated_stem_rich,
        "annotatedOptions": annotated_options,
        "annotatedOptionRich": annotated_option_rich,
    }


def analysis_streams(document: DocumentObject) -> list[list[Paragraph]]:
    streams = [all_body_paragraphs(document), document.paragraphs, leaf_body_paragraphs(document)]
    streams.extend(textbox_streams(document))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.paragraphs:
                    streams.append(cell.paragraphs)
    return streams


def similarity(left: str, right: str) -> float:
    left_key = match_text(left)[:800]
    right_key = match_text(right)[:800]
    if not left_key or not right_key:
        return 0.0
    return SequenceMatcher(None, left_key, right_key, autojunk=False).ratio()


def merge_analysis(path: Path, bases: list[dict[str, object]]) -> None:
    document = Document(path)
    primary_chunks = split_numbered_chunks(all_body_paragraphs(document), analysis=True)
    if len(primary_chunks) != 300:
        raise ValueError(f"{path.name}: expected 300 ordered analysis questions, found {len(primary_chunks)}")
    primary = [parse_analysis_candidate(chunk) for chunk in primary_chunks]
    if any(candidate is None for candidate in primary):
        missing = [index + 1 for index, candidate in enumerate(primary) if candidate is None]
        raise ValueError(f"{path.name}: failed to parse ordered analysis questions {missing}")

    candidates: list[dict[str, object]] = [candidate for candidate in primary if candidate]
    for stream in analysis_streams(document):
        for chunk in split_numbered_chunks(stream, analysis=True):
            candidate = parse_analysis_candidate(chunk)
            if candidate:
                candidates.append(candidate)

    by_source: dict[str, list[dict[str, object]]] = defaultdict(list)
    for candidate in candidates:
        by_source[str(candidate.get("sourceLabel") or "")].append(candidate)

    for index, (base, ordered) in enumerate(zip(bases, primary)):
        assert ordered is not None
        expected_local = index % 20 + 1
        if int(ordered["localNumber"]) != expected_local:
            raise ValueError(f"{path.name}: ordered analysis mismatch at global item {index + 1}")
        source_choices = by_source[str(ordered.get("sourceLabel") or "")] or [ordered]
        ranked = sorted(
            ((similarity(str(base["stem"]), str(candidate["annotatedStem"])), candidate) for candidate in source_choices),
            key=lambda item: item[0],
            reverse=True,
        )
        best_score, best_rich = ranked[0]
        merged = dict(ordered)
        if best_score >= 0.48:
            for field in ("annotatedStem", "annotatedStemRich", "annotatedOptions", "annotatedOptionRich"):
                merged[field] = best_rich.get(field)
            pairing_mode = "text-and-sequence"
        else:
            # A handful of pages store the annotated passage as a positioned drawing. The clean
            # book text remains authoritative while source/answer/analysis stay paired by the
            # already-validated 1..20 sequence within the set.
            merged["annotatedStem"] = base["stem"]
            merged["annotatedStemRich"] = base["stemRich"]
            merged["annotatedOptions"] = base["options"]
            merged["annotatedOptionRich"] = base["optionRich"]
            pairing_mode = "validated-sequence-fallback"

        similar_choices = [candidate for score, candidate in ranked if score >= best_score - 0.08]
        for field in ("answer", "typeAndPassage", "practicalAnalysis", "source", "sourceLabel", "accuracy"):
            if not merged.get(field):
                merged[field] = next((candidate.get(field) for candidate in similar_choices if candidate.get(field)), None)
        base.update(merged)
        option_sources = [base.get("options", [])] + [candidate.get("annotatedOptions", []) for candidate in [ordered, best_rich, *source_choices]]
        option_values: dict[str, str] = {}
        for source_options in option_sources:
            for option in source_options or []:
                label = str(option.get("label"))
                content = str(option.get("content") or "").strip()
                if label not in {"A", "B", "C", "D"} or not content or label in option_values:
                    continue
                content = re.sub(r"[①-⑳㉑-㊿]+$", "", content).strip()
                content = re.sub(r"\s*[（(](?:片面|无中生有|具体类|出处有误|表述有误|强加因果|区别比较联系类|中性表达)[^）)]*[）)]\s*$", "", content)
                option_values[label] = content
        option_values.update(KNOWN_OPTION_REPAIRS.get((int(base["setNumber"]), int(base["localNumber"])), {}))
        base["options"] = [{"label": label, "content": option_values[label]} for label in "ABCD" if label in option_values]
        base["optionRich"] = {label: [{"text": option_values[label]}] for label in option_values}

        if not base.get("answer"):
            practical = str(base.get("practicalAnalysis") or "")
            inferred = re.search(r"(?:锁定|答案为|对应)\s*([A-D])\s*选项|([A-D])\s*选项[^。；\n]{0,40}当选", practical)
            if inferred:
                base["answer"] = (inferred.group(1) or inferred.group(2)).upper()
        if not base.get("typeAndPassage"):
            stem = str(base.get("stem") or "")
            if "重新排列" in stem:
                inferred_type = "语句排序类"
            elif "填入文中画横线" in stem or "填入横线" in stem:
                inferred_type = "语句填入类"
            elif "接下来" in stem:
                inferred_type = "下文推断类"
            elif any(value in stem for value in ("说法正确", "说法错误", "相符的是", "不相符的是")):
                inferred_type = "细节判断类"
            else:
                inferred_type = "中心理解题"
            base["typeAndPassage"] = inferred_type + "+科普介绍文章"
        base["matchScore"] = round(max(best_score, similarity(str(base["stem"]), str(ordered["annotatedStem"]))), 4)
        base["pairingMode"] = pairing_mode


def marker_token(text: str) -> tuple[str, int | None] | None:
    position = text.find(NOTE_LABEL)
    if position < 0:
        return None
    prefix = text[:position].rstrip()
    index = len(prefix) - 1
    while index >= 0 and (prefix[index] in MARKER_CHARS or prefix[index].isspace()) and len(prefix) - index <= 8:
        index -= 1
    token = re.sub(r"\s+", "", prefix[index + 1 :])
    if not token:
        return None

    values: list[int] = []
    for char in token:
        if char in CIRCLED_VALUE:
            values.append(CIRCLED_VALUE[char])
        elif char.isdigit():
            values.append(int(char))
    if not values:
        return token, None
    if len(values) == 1:
        number = values[0]
    else:
        number = int("".join(str(value % 10) for value in values))
    return token, number


def extract_notes(document: DocumentObject) -> list[list[dict[str, object]]]:
    paragraphs = leaf_body_paragraphs(document)
    records: list[dict[str, object]] = []
    current: dict[str, object] | None = None

    def finish() -> None:
        nonlocal current
        if not current:
            return
        body = [compact_text(value) for value in current.pop("parts") if compact_text(value) and not is_noise(value)]
        current["content"] = "\n".join(body).strip()
        records.append(current)
        current = None

    boundaries = (ANSWER_LABEL, PRACTICAL_LABEL, "【题型")
    for paragraph in paragraphs:
        text = compact_text(paragraph.text)
        if not text:
            continue
        if NOTE_LABEL in text:
            finish()
            token_info = marker_token(text)
            if not token_info:
                continue
            token, number = token_info
            suffix = text.split(NOTE_LABEL, 1)[1].lstrip("：: ")
            current = {"marker": token, "number": number, "parts": [suffix] if suffix else []}
            continue
        if current:
            if ANALYSIS_START.match(text) or any(boundary in text for boundary in boundaries):
                finish()
                continue
            current["parts"].append(text)
    finish()

    groups: list[list[dict[str, object]]] = []
    current_group: list[dict[str, object]] = []
    for record in records:
        if record.get("number") == 1 and current_group:
            groups.append(current_group)
            current_group = []
        current_group.append(record)
    if current_group:
        groups.append(current_group)
    # Two lower-volume set openings use a broken glyph for ①. Recover those boundaries
    # from the sharp high-to-low marker reset while keeping realistic per-set note counts.
    while len(groups) < 15:
        best: tuple[float, int, int] | None = None
        for group_index, group in enumerate(groups):
            if len(group) < 45:
                continue
            for cut in range(15, len(group) - 14):
                previous = group[cut - 1].get("number")
                following = group[cut].get("number")
                if not isinstance(previous, int) or not isinstance(following, int):
                    continue
                if previous < 18 or following > 6:
                    continue
                score = abs(cut - 35) + abs((len(group) - cut) - 35)
                if best is None or score < best[0]:
                    best = (score, group_index, cut)
        if best is None:
            break
        _, group_index, cut = best
        group = groups[group_index]
        groups[group_index : group_index + 1] = [group[:cut], group[cut:]]
    return groups


def extract_inline_markers(text: str) -> set[int]:
    numbers: set[int] = set()
    for char in text:
        if char in CIRCLED_VALUE:
            numbers.add(CIRCLED_VALUE[char])
    return numbers


def note_similarity(note: str, question: dict[str, object]) -> float:
    haystack = " ".join(
        str(question.get(field) or "")
        for field in ("annotatedStem", "practicalAnalysis", "stem", "typeAndPassage")
    )
    left = match_text(note)
    right = match_text(haystack)
    if not left or not right:
        return 0.0
    sample = left[:240]
    best = SequenceMatcher(None, sample, right, autojunk=False).ratio()
    # Shared CJK bigrams make topic-specific notes rank much more reliably than generic prose.
    left_pairs = {sample[index : index + 2] for index in range(max(0, len(sample) - 1))}
    right_pairs = {right[index : index + 2] for index in range(max(0, len(right) - 1))}
    if left_pairs:
        best += len(left_pairs & right_pairs) / len(left_pairs)
    return best


def attach_notes(path: Path, bases: list[dict[str, object]]) -> tuple[int, int]:
    groups = extract_notes(Document(path))
    if len(groups) != 15:
        raise ValueError(f"{path.name}: expected 15 note groups, found {len(groups)}")
    attached = 0
    unmatched = 0
    for set_offset, notes in enumerate(groups):
        questions = bases[set_offset * 20 : (set_offset + 1) * 20]
        for question in questions:
            combined = " ".join(
                [str(question.get("annotatedStem") or ""), str(question.get("practicalAnalysis") or "")]
                + [str(option.get("content") or "") for option in question.get("annotatedOptions", [])]
            )
            question["inlineMarkers"] = sorted(extract_inline_markers(combined))
            question["notes"] = []

        for position, note in enumerate(notes):
            if not str(note.get("content") or "").strip():
                continue
            number = note.get("number")
            eligible = [question for question in questions if number in question.get("inlineMarkers", [])]
            all_ranked = sorted(
                ((note_similarity(str(note["content"]), question), question) for question in questions),
                key=lambda item: item[0],
                reverse=True,
            )
            eligible_ranked = sorted(
                ((note_similarity(str(note["content"]), question), question) for question in eligible),
                key=lambda item: item[0],
                reverse=True,
            )
            # A marker is only a candidate hint. If the best content match is
            # materially stronger than the marker-selected question, prefer
            # content so misplaced/duplicated marker numbers cannot scramble
            # annotations.
            if eligible_ranked and eligible_ranked[0][0] >= all_ranked[0][0] - 0.10:
                score, selected = eligible_ranked[0]
                pairing_mode = "marker-and-content"
            else:
                score, selected = all_ranked[0]
                pairing_mode = "content-similarity"
            if score < 0.10 and not eligible:
                unmatched += 1
            selected["notes"].append(
                {
                    "marker": note["marker"],
                    "number": note.get("number"),
                    "content": note["content"],
                    "matchScore": round(score, 4),
                    "order": position,
                    "pairingMode": pairing_mode,
                }
            )
            attached += 1
    return attached, unmatched


def sql_quote(value: str) -> str:
    return "'" + value.replace("'", "''") + "'"


def difficulty(accuracy: int | None) -> str:
    if accuracy is None or accuracy >= 70:
        return "基础"
    if accuracy >= 45:
        return "进阶"
    return "挑战"


def question_type(type_and_passage: str | None) -> str:
    if not type_and_passage:
        return "片段阅读"
    return re.split(r"[+＋]", type_and_passage, maxsplit=1)[0].strip() or "片段阅读"


def build_record(question: dict[str, object], global_number: int) -> dict[str, object]:
    practical = str(question.get("practicalAnalysis") or "").strip()
    notes = question.get("notes", [])
    explanation_parts = [practical] if practical else []
    if notes:
        explanation_parts.append(
            "花生批注：\n"
            + "\n".join(f"{note['marker']} {note['content']}" for note in notes)
        )
    details = {
        "globalNumber": global_number,
        "setNumber": question["setNumber"],
        "localNumber": question["localNumber"],
        "sourceLabel": question.get("sourceLabel"),
        "accuracy": question.get("accuracy"),
        "typeAndPassage": question.get("typeAndPassage"),
        "stemRich": question["stemRich"],
        "optionRich": question["optionRich"],
        "annotatedStemRich": question.get("annotatedStemRich"),
        "annotatedOptionRich": question.get("annotatedOptionRich"),
        "practicalAnalysis": practical,
        "notes": [
            {
                "marker": note["marker"],
                "content": note["content"],
                "matchScore": note.get("matchScore"),
                "pairingMode": note.get("pairingMode"),
                "order": note.get("order"),
            }
            for note in notes
        ],
        "pairingMode": question.get("pairingMode"),
    }
    return {
        "id": 300000 + global_number,
        "categoryId": 3,
        "type": question_type(question.get("typeAndPassage")),
        "stem": question["stem"],
        "options": question["options"],
        "answer": question.get("answer"),
        "explanation": "\n\n".join(explanation_parts),
        "source": question.get("source") or "片段阅读600题",
        "difficulty": difficulty(question.get("accuracy")),
        "status": "published",
        "details": details,
        "matchScore": question.get("matchScore"),
        "noteCount": len(notes),
    }


def validate(records: list[dict[str, object]]) -> dict[str, object]:
    issues: list[str] = []
    if len(records) != 600:
        issues.append(f"question count is {len(records)}, expected 600")
    ids = [record["id"] for record in records]
    if len(set(ids)) != len(ids):
        issues.append("duplicate ids")
    for index, record in enumerate(records, 1):
        if record["answer"] not in {"A", "B", "C", "D"}:
            issues.append(f"question {index}: missing/invalid answer")
        labels = [option["label"] for option in record["options"]]
        if labels != ["A", "B", "C", "D"]:
            issues.append(f"question {index}: invalid options {labels}")
        details = record["details"]
        if not details.get("typeAndPassage"):
            issues.append(f"question {index}: missing type and passage")
        if not details.get("practicalAnalysis"):
            issues.append(f"question {index}: missing practical analysis")
        if not record["explanation"]:
            issues.append(f"question {index}: missing explanation")
        if float(record.get("matchScore") or 0) < 0.48 and details.get("pairingMode") != "validated-sequence-fallback":
            issues.append(f"question {index}: low book-analysis match")
    return {
        "questionCount": len(records),
        "answerCount": sum(bool(record.get("answer")) for record in records),
        "typeCount": sum(bool(record["details"].get("typeAndPassage")) for record in records),
        "practicalAnalysisCount": sum(bool(record["details"].get("practicalAnalysis")) for record in records),
        "questionsWithNotes": sum(bool(record["details"].get("notes")) for record in records),
        "noteCount": sum(len(record["details"].get("notes", [])) for record in records),
        "underlinedQuestionCount": sum(
            any(run.get("underline") for run in record["details"].get("stemRich", []))
            for record in records
        ),
        "minimumMatchScore": min(float(record.get("matchScore") or 0) for record in records),
        "sequenceFallbackCount": sum(record["details"].get("pairingMode") == "validated-sequence-fallback" for record in records),
        "issues": issues,
    }


def write_sql(records: list[dict[str, object]], output: Path) -> None:
    lines = [
        "-- Refresh the 600 passage-reading questions while preserving their IDs.",
        "-- Preserving IDs keeps existing attempts and the wrongbook linked to the refreshed questions.",
        "DELETE FROM questions WHERE category_id = 3;",
        "",
    ]
    for record in records:
        values = [
            str(record["id"]),
            str(record["categoryId"]),
            sql_quote(str(record["type"])),
            sql_quote(str(record["stem"])),
            sql_quote(json.dumps(record["options"], ensure_ascii=False, separators=(",", ":"))),
            sql_quote(str(record["answer"])),
            sql_quote(str(record["explanation"])),
            sql_quote(str(record["source"])),
            sql_quote(str(record["difficulty"])),
            sql_quote(str(record["status"])),
            sql_quote(json.dumps(record["details"], ensure_ascii=False, separators=(",", ":"))),
        ]
        lines.append(
            "INSERT INTO questions (id, category_id, type, stem, options_json, answer, explanation, source, difficulty, status, details_json) VALUES ("
            + ",".join(values)
            + ");"
        )
    output.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    parser = argparse.ArgumentParser(description="Pair and import the 600 passage-reading questions from four Word documents.")
    default_desktop = Path.home() / "Desktop" / "三板块"
    parser.add_argument("--desktop", type=Path, default=default_desktop if default_desktop.exists() else Path.home() / "Desktop")
    parser.add_argument("--output-json", type=Path, default=Path("generated/reading-600.json"))
    parser.add_argument("--output-sql", type=Path, default=Path("migrations/0005_replace_with_reading_600.sql"))
    parser.add_argument("--report", type=Path, default=Path("generated/reading-600-report.json"))
    args = parser.parse_args()

    desktop = args.desktop
    def find_document(kind: str, volume: str) -> Path:
        suffix = f"{kind}【{volume}】.docx"
        matches = sorted(path for path in desktop.iterdir() if path.is_file() and path.name.endswith(suffix))
        if not matches:
            raise FileNotFoundError(f"找不到{kind}{volume}册 DOCX")
        return matches[0]

    book_up = find_document("题本", "上")
    book_down = find_document("题本", "下")
    analysis_up = find_document("解析", "上")
    analysis_down = find_document("解析", "下")

    upper = parse_book(book_up, 1)
    lower = parse_book(book_down, 16)
    merge_analysis(analysis_up, upper)
    merge_analysis(analysis_down, lower)
    upper_note_count, upper_unmatched = attach_notes(analysis_up, upper)
    lower_note_count, lower_unmatched = attach_notes(analysis_down, lower)

    records = [build_record(question, index) for index, question in enumerate(upper + lower, 1)]
    report = validate(records)
    report.update(
        {
            "sourceDocuments": [book_up.name, book_down.name, analysis_up.name, analysis_down.name],
            "extractedNotes": upper_note_count + lower_note_count,
            "lowConfidenceNoteAssignments": upper_unmatched + lower_unmatched,
        }
    )
    if report["issues"]:
        raise ValueError("validation failed:\n" + "\n".join(report["issues"][:80]))

    args.output_json.parent.mkdir(parents=True, exist_ok=True)
    args.output_sql.parent.mkdir(parents=True, exist_ok=True)
    args.report.parent.mkdir(parents=True, exist_ok=True)
    args.output_json.write_text(json.dumps(records, ensure_ascii=False, indent=2), encoding="utf-8")
    write_sql(records, args.output_sql)
    args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
